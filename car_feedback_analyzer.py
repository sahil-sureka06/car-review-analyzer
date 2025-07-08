import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import plotly.express as px
import re
from collections import defaultdict
from ibm_watson import NaturalLanguageUnderstandingV1
from ibm_cloud_sdk_core.authenticators import IAMAuthenticator
from ibm_watson.natural_language_understanding_v1 import Features, SentimentOptions

# ───────────────────────────── Streamlit page ─────────────────────────────
st.set_page_config(page_title="🚗 Car Review Analyzer - Sahil Sureka", layout="wide")

# ───────────────────────────── Sidebar ─────────────────────────────
st.sidebar.title("📄 Project Details")
st.sidebar.markdown(
    """
**Car Feedback Analyzer**  
Get AI-powered sentiment analysis & issue detection for car rental feedback.  
**Made by Sahil Sureka** | Powered by IBM Watson NLU
"""
)
st.sidebar.markdown("🔗 [GitHub Repo](https://github.com/sahil-sureka06/car-review-analyzer)")

# ───────────────────────────── Custom CSS ─────────────────────────────
st.markdown(
    """
    <style>
    .main-header {
        text-align: center;
        padding: 1.5rem;
        background: linear-gradient(135deg, #6A0572 0%, #AB83A1 100%);
        border-radius: 12px;
        color: white;
        margin-bottom: 2rem;
    }
    .upload-section {
        border: 2px dashed #6A0572;
        border-radius: 12px;
        padding: 2rem;
        background: #f9f0fb;
        text-align: center;
        margin-bottom: 1rem;
    }
    .summary-card {
        background: #ffffff;
        padding: 1rem;
        border-radius: 10px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        text-align: center;
        margin-bottom: 1rem;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ───────────────────────────── IBM Watson init ─────────────────────────────
@st.cache_resource
def init_watson():
    api_key = "Ymi5EPk-qULX8RUiLFKKDvTlKRtl4CMurQKmOuYS0aQm"
    service_url = "https://api.au-syd.natural-language-understanding.watson.cloud.ibm.com/instances/c8ef4303-e705-4357-a9f4-d63007836993"

    authenticator = IAMAuthenticator(api_key)
    nlu = NaturalLanguageUnderstandingV1(
        version="2023-06-01",
        authenticator=authenticator,
    )
    nlu.set_service_url(service_url)
    return nlu

# ───────────────────────────── Keyword map ─────────────────────────────
ISSUE_KEYWORDS = {
    "Service Problem": r"(issue|problem|breakdown|engine|mechanic|maintenance)",
    "Cleanliness": r"(dirty|unclean|smell|stain|clean)",
    "Delay": r"(slow|wait|delay|late|queue|pickup)",
    "Staff Behavior": r"(rude|unhelpful|impolite|staff|service)",
    "Pricing Issue": r"(expensive|overpriced|hidden fee|charge)",
    "Tech Issues": r"(gps|navigation|bluetooth|usb)",
    "Damage": r"(scratch|dent|damage)"
}

# ───────────────────────────── Sentiment Analysis ─────────────────────────────
@st.cache_data(show_spinner=False)
def watson_sentiment(text: str):
    nlu = init_watson()
    resp = nlu.analyze(text=text, features=Features(sentiment=SentimentOptions())).get_result()
    label = resp["sentiment"]["document"]["label"].lower()
    score = resp["sentiment"]["document"]["score"]
    return label, score

def analyze_review(text: str):
    sentiment, _ = watson_sentiment(text)
    issues = [label for label, pattern in ISSUE_KEYWORDS.items() if re.search(pattern, text, re.I)]
    return sentiment, sorted(set(issues))

# ───────────────────────────── Generate Word Report ─────────────────────────────
def generate_word_report(results):
    doc = Document()
    doc.add_heading('Car Review Analysis Report', 0)
    doc.add_paragraph('Prepared by: Sahil Sureka')

    for idx, item in enumerate(results, 1):
        doc.add_heading(f'Review {idx}', level=1)
        doc.add_paragraph(f"Feedback: {item['review']}")
        doc.add_paragraph(f"Sentiment: {item['sentiment'].capitalize()}")
        doc.add_paragraph(f"Rating: {item['rating']}")
        doc.add_paragraph(f"Issues Detected: {', '.join(item['issues']) if item['issues'] else 'None'}")
        doc.add_paragraph('-' * 40)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ───────────────────────────── UI Layout ─────────────────────────────
st.markdown('<div class="main-header"><h1>🚗 Car Feedback Analyzer</h1><p>Watson AI-Powered Sentiment & Issue Detection</p></div>', unsafe_allow_html=True)

st.markdown("### 📥 Upload Customer Feedback")
uploaded_file = st.file_uploader("Upload CSV file", type="csv")
use_demo = st.checkbox("Use Demo Data Instead")

if uploaded_file:
    df = pd.read_csv(uploaded_file)
elif use_demo:
    df = pd.DataFrame([
        {"customer_id": 1, "review": "The car was clean but late.", "rating": 3},
        {"customer_id": 2, "review": "Excellent service and very helpful staff.", "rating": 5},
        {"customer_id": 3, "review": "Vehicle had a GPS issue.", "rating": 2},
        {"customer_id": 4, "review": "Smooth ride, no complaints.", "rating": 4},
        {"customer_id": 5, "review": "Hidden charges and rude staff.", "rating": 1},
    ])
else:
    df = None

if df is not None:
    st.write("### Preview Data")
    st.dataframe(df)

    if st.button("🔍 Run AI Analysis"):
        with st.spinner("Analyzing with IBM Watson..."):
            results = []
            issue_counter = defaultdict(int)

            for _, row in df.iterrows():
                sentiment, issues = analyze_review(row['review'])
                results.append({"review": row['review'], "sentiment": sentiment, "issues": issues, "rating": row.get('rating', 'N/A')})
                for issue in issues:
                    issue_counter[issue] += 1

        st.success("Analysis Completed!")

        sentiments = pd.Series([r['sentiment'] for r in results]).value_counts()
        st.write("### 📊 Sentiment Distribution")
        st.plotly_chart(px.pie(values=sentiments.values, names=sentiments.index, title="Sentiment Breakdown"), use_container_width=True)

        if issue_counter:
            st.write("### 🚩 Common Issues")
            st.dataframe(pd.DataFrame({"Issue": list(issue_counter.keys()), "Count": list(issue_counter.values())}).sort_values("Count", ascending=False))

        word_buffer = generate_word_report(results)
        st.download_button(
            label="📥 Download Word Report",
            data=word_buffer,
            file_name="Car_Review_Analysis_Report_Sahil_Sureka.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

st.markdown("---")
st.markdown("<div style='text-align:center; color:gray'>Made with ❤️ by <strong>Sahil Sureka</strong></div>", unsafe_allow_html=True)
